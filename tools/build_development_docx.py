from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "DEVELOPMENT_DOCUMENT.md"
OUTPUT = ROOT / "Talent_Match_Development_Document.docx"

BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(85, 85, 85)
BORDER = "DADCE0"
HEADER_FILL = "F5F5F5"


def set_run_font(run, name="Arial", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=8, line_spacing=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def set_cell_text(cell, text, bold=False, fill=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0)
    run = p.add_run(text)
    set_run_font(run, size=9.5, bold=bold)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), BORDER)
        borders.append(tag)
    tbl_pr.append(borders)


def add_markdown_table(doc, rows):
    cleaned = []
    for row in rows:
        cells = [cell.strip() for cell in row.strip().strip("|").split("|")]
        if all(set(cell) <= {"-", ":", " "} for cell in cells):
            continue
        cleaned.append(cells)
    if not cleaned:
        return
    width = max(len(row) for row in cleaned)
    table = doc.add_table(rows=len(cleaned), cols=width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for r_idx, row in enumerate(cleaned):
        for c_idx in range(width):
            text = row[c_idx] if c_idx < len(row) else ""
            set_cell_text(table.cell(r_idx, c_idx), text, bold=(r_idx == 0), fill=HEADER_FILL if r_idx == 0 else None)
    doc.add_paragraph()


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=8, line_spacing=1.0)
    run = p.add_run("\n".join(lines))
    set_run_font(run, name="Courier New", size=8.5, color=RGBColor(40, 40, 40))


def add_inline_runs(paragraph, text, size=11):
    parts = text.split("`")
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        if idx % 2 == 1:
            set_run_font(run, name="Courier New", size=size - 0.5, color=RGBColor(40, 40, 40))
        else:
            set_run_font(run, size=size)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)

    for name, size, color in [
        ("Heading 1", 20, BLACK),
        ("Heading 2", 16, BLACK),
        ("Heading 3", 14, RGBColor(67, 67, 67)),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = color

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.add_run("Talent Match Development Document"), size=9, color=MUTED)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    table_buffer = []
    code_buffer = []
    in_code = False

    for line in lines:
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_buffer)
                code_buffer = []
                in_code = False
            else:
                if table_buffer:
                    add_markdown_table(doc, table_buffer)
                    table_buffer = []
                in_code = True
            continue

        if in_code:
            code_buffer.append(line)
            continue

        if line.strip().startswith("|"):
            table_buffer.append(line)
            continue

        if table_buffer:
            add_markdown_table(doc, table_buffer)
            table_buffer = []

        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, after=3)
            run = p.add_run(stripped[2:])
            set_run_font(run, size=26)
        elif stripped.startswith("## "):
            p = doc.add_paragraph(stripped[3:], style="Heading 1")
            set_paragraph_spacing(p, before=20, after=6)
        elif stripped.startswith("### "):
            p = doc.add_paragraph(stripped[4:], style="Heading 2")
            set_paragraph_spacing(p, before=18, after=6)
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            set_paragraph_spacing(p, after=4)
            add_inline_runs(p, stripped[2:])
        elif stripped[0:3] in [f"{i}. " for i in range(1, 10)]:
            p = doc.add_paragraph(style="List Number")
            set_paragraph_spacing(p, after=4)
            add_inline_runs(p, stripped[3:])
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=4, after=8)
            run = p.add_run(stripped[2:])
            set_run_font(run, size=10.5, color=MUTED, italic=True)
        else:
            p = doc.add_paragraph()
            set_paragraph_spacing(p)
            add_inline_runs(p, stripped)

    if table_buffer:
        add_markdown_table(doc, table_buffer)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
